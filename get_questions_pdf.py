import docx
from pathlib import Path
import re
from collections import deque
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pymupdf

import tkinter as tk
from tkinter import messagebox
from UI.ui import manage_ui
from UI.ui import UIContext


'''
TODO: Extra Lernziele 7 - question 11 and 7: all lines concatenated in the same line
TODO: when adding "split" to solutions some lines concatenate without spaces
TODO: set multiline questions to bold (question 14 Bewegungsaparat)
TODO: set numbered lists when required
TODO: check question 28 the text at the end that is not part of the bullets is added as a bullet
'''

class ParserState:
    solution_page: int
    solution_block: int
    solution_line: int
    solution_span: int

    def __init__(self):
        self.solution_page = 0
        self.solution_block = 0
        self.solution_line = 0
        self.solution_span = 0


def starts_with_number(line):
    return re.match(r'[0-9]+\)', line)

def remove_number(line):
    # sub eliminates the string when it finds it
    # first with ^ it forces to check it only at the start
    # Then it can have more than one number and then parenthesis and spaces
    return(re.sub(r'^[0-9]+\)\s*', '', line,))

def is_questions_page(line):
    return line.startswith("Lernziele")

def is_solutions_page(line):
    return line.startswith("Lösungen")

def can_be_omitted(line):
    return ("Medizinische Grundlagen; 2026-2027" in line or
            "Barbara" in line or
            re.match(r'^[0-9]\s', line) or
            line == " " or
            "Lösungen" in line
            )

def is_new_line(line: str):
    return (re.match(r'^[0-9].\s', line) or
            line.startswith("-")
            )

def is_bold(flags: int) -> bool:
    return (flags & 2 ** 4) > 0

def is_italic(flags: int) -> bool:
    return (flags & 2 ** 1) > 0

def is_underlined(flags: int) -> bool:
    return (flags & 2 ** 0) > 0

def is_bullet_span(span):
    text: str
    text = span["text"][0]
    bullet_list = {"•", "◦", "▪", "-", "-", "*"}
    return text in bullet_list
    # bullet_list = {"•", "◦", "▪", "-", "-", "*"}
    # return text in bullet_list

def remove_bullet(span) -> str:
    text: str
    text = span["text"][1:]
    return text.lstrip()


def add_run_with_format(paragraph, text: str, span_flags: int) -> None:
    run = paragraph.add_run(text)
    run.bold = is_bold(span_flags)
    run.italic = is_italic(span_flags)
                   

def get_solutions(output_file: docx.Document, solutions_full_text: list, parser_state: ParserState) -> None:

    line: str
    added_paragraph = None
    is_new_solution = False             # If this is true, this is a new solution and the previous has ended
    solution_ended = False              # If this is true, the current solution has been fully parsed
    page_num, block_num, line_num, span_num = parser_state.solution_page, parser_state.solution_block, parser_state.solution_line, parser_state.solution_span

    # for page_num, page in enumerate(solutions_full_text):
    #     page = solutions_full_text[page_num]
    #     text_blocks = page.get_text("dict", flags=pymupdf.TEXTFLAGS_TEXT)["blocks"]
    #     for block_num, block in enumerate(text_blocks):
    #         lines = len(block["lines"])
    #         print(f"Page {page_num} - Block {block_num} - {lines} ")

    # pass

    while page_num < len(solutions_full_text) and not(solution_ended):     
        page = solutions_full_text[page_num]
        text_blocks = page.get_text("dict", flags=pymupdf.TEXTFLAGS_TEXT)["blocks"]
        
        while block_num < len(text_blocks) and not(solution_ended):
            block = text_blocks[block_num]
            
            while line_num < len(block["lines"]) and not(solution_ended):
                line = block["lines"][line_num]

                while span_num < len(line["spans"]):
                    span = line["spans"][span_num]
                    span_text = span["text"]
                    span_flags = span["flags"]

                    if span_text.startswith("1)"):
                        pass

                    if span_num == 0:
                        if can_be_omitted(span_text):
                            span_num += 1
                            continue

                        if starts_with_number(span_text):                           
                            # If the line starts with number and it's a new solution, the entire solution was already added
                            if is_new_solution:
                                solution_ended = True
                                break

                            span_text = remove_number(span_text)
                            is_new_solution = True

                            # For cases where the question starts with "10) "(Like L 7- F 10) to not create extra lines
                            if len(span_text) > 0:
                                added_paragraph = output_file.add_paragraph()

                        elif is_bullet_span(span):        # If the span starts with - then add it as a bullet list and remove the text
                            span_text = remove_bullet(span)
                            added_paragraph = output_file.add_paragraph(style="List Bullet")
                        elif is_bold(span_flags):           # If the span starts with bold, then it is a new paragraph in the same solution
                            added_paragraph = output_file.add_paragraph()
                    else:
                        if not(added_paragraph):
                            added_paragraph = output_file.add_paragraph()


                    # Add a run only if the current span is not a heading (headings have a size of 14 or more)                    
                    if added_paragraph and span["size"] < 14:
                        add_run_with_format(added_paragraph, span_text, span_flags)
                    span_num += 1
                
                if not(solution_ended):
                    span_num = 0
                    line_num += 1

            if not(solution_ended):        
                line_num = 0
                block_num += 1       
        
        if not(solution_ended):
            block_num = 0
            page_num += 1
                 

    added_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    added_paragraph.paragraph_format.space_after = 0
    
    # Store the state of the parser to start from this section
    parser_state.solution_page, parser_state.solution_block, parser_state.solution_line, parser_state.solution_span = page_num, block_num, line_num, span_num

    # Add an additional line
    output_file.add_paragraph("")

def create_questions_file(ui_ctx: UIContext):

    # --------------------------------------------
    # Open the file to read
    # --------------------------------------------

    try:
        file_path = Path(ui_ctx._input_file_name).absolute()
    except (UnboundLocalError, TypeError):
        tk.messagebox.showerror(title="File not found", message=f"Select a .pdf file.")
        
    try:
        questions_file = pymupdf.open(file_path)
    except FileNotFoundError:
        print(f"File {file_path} not found.")
        tk.messagebox.showerror(title="File not found", message=f"File {file_path} not found.")
        # sys.exit(1)


    # --------------------------------------------
    # Open the file to write the final results
    # --------------------------------------------

    if ui_ctx.output_file_name == None:
        tk.messagebox.showerror(title="File not found", message=f"Select an output folder.")

    output_file_name = file_path.stem
    output_file_path = Path(ui_ctx.output_file_name) / Path(output_file_name + '.docx')
    output_file = docx.Document()

    try:
        with open(output_file_path, "a"):
            pass
    except PermissionError:
        print(f"File {file_path} is locked or already open.")
        tk.messagebox.showerror(title="File locked", message=f"File {file_path} is locked or already open.")
        # sys.exit(1)


    # --------------------------------------------
    # Look for the solutions page
    # --------------------------------------------

    solutions_page = None
    for page_num, page in enumerate(questions_file):
        lines = page.get_text("text").splitlines()
        for line in lines:
            if starts_with_number(line) or is_questions_page(line):
                print(f"{page} not solutions page")
                break
            if is_solutions_page(line):
                solutions_page = page_num
                break
        if solutions_page is not None:
            break

    # if solutions_page is None:
    #     raise ValueError("Solutions page not found.")
    

    # --------------------------------------------
    # Create a list with the full text of the questions and solutions
    # --------------------------------------------

    questions_full_text = questions_file[0:solutions_page]
    solutions_full_text = questions_file[solutions_page:]


    # --------------------------------------------
    # Start saving the contents in a new file
    # --------------------------------------------

    is_new_question = False                     # To detect if the previous line was a question so that it is not formatted as a headline
    questions_initiated = False             # To detect once the first lines of the document are omitted
    question_number = 0
    prev_heading = False                    # If the previous line was a heading, do not add answer
    is_bullet = False                       # For questions with bullets
    parser_state = ParserState()

    for page in questions_full_text:
        text_blocks = page.get_text("dict", flags=pymupdf.TEXTFLAGS_TEXT)["blocks"]
        
        for block_num, block in enumerate(text_blocks):
            for line_num, line in enumerate(block["lines"]):
                for span_num, span in enumerate(line["spans"]):
                    
                    span_text = span["text"]
                    span_flags = span["flags"]

                    # Check this for the first words from each line
                    if span_num == 0:
                        # If line is not questions (the same as is not a questions page) go to the next line
                        if not(is_questions_page(span_text)) and not(questions_initiated):
                            continue
                        elif not(questions_initiated):
                            questions_initiated = True
                                    
                        # If the current line is a question
                        if starts_with_number(span_text):
                            is_new_question = True
                            question_number += 1                       
                        else:
                            if can_be_omitted(span_text):
                                continue

                            # If it is not a question, then it is a heading so we have to add a new paragraph
                            # Otherwise the previous question continues and we don't have to create a new paragraph.
                            if not(is_bold(span_flags)):
                                is_new_question = False         # This is because this is the continuation of a previous question
                                if is_bullet_span(span):        # If the span starts with - then add it as a bullet list and remove the text
                                    is_bullet = True
                                    span_text = ""
                            else:
                                is_heading = True

                    else:
                        add_run_with_format(added_paragraph, span_text.strip(), span_flags)

                    
                    # If this is a new question, add the solution from the previous question
                    if question_number > 1 and is_new_question and not(prev_heading):
                        prev_heading = False
                        
                        added_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        # TODO: set the question to bold
                        for run in added_paragraph.runs:
                            run.bold = True

                        # Add the solution and an additional line
                        get_solutions(output_file, solutions_full_text, parser_state)


                    # The paragraph must be added at the end to add the solution of the previous question
                    # If this is a new question or a heading, add a new paragraph
                    if is_new_question or is_heading:
                        added_paragraph = output_file.add_paragraph()
                    
                    # In any case, add the text as a run in the current paragraph, using strip to remove trailing spaces
                    add_run_with_format(added_paragraph, span_text.strip(), span_flags)

                    # If it is heading, set the style as a heading
                    if is_heading:
                        added_paragraph.style = "Heading 2"
                        added_paragraph.paragraph_format.space_after = Pt(12)
                        is_heading = False
                        prev_heading = True
                    else:
                        prev_heading = False
                        if is_bullet:
                            added_paragraph = output_file.add_paragraph(style="List Bullet")
                            is_bullet = False

    # Format last question
    for run in added_paragraph.runs:
        run.bold = True
    
    # Add the solution for the final question and an additional line
    get_solutions(output_file, solutions_full_text, parser_state)

    # Change the first paragraph as a title
    output_file.paragraphs[0].style = "Title"
    output_file.paragraphs[0].style.font.size = Pt(20)

    # Chenge the font to calibri for normal
    style = output_file.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    output_file.save(output_file_path)
    tk.messagebox.showinfo(title="File created", message=f"File \"{output_file_name}.docx\" successfully created!")


# TODO: Delete this later, just for testing
ui_ctx = UIContext()
# ui_ctx.input_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py\\2. Lernziele Infektionslehre und Epidemiologie.pdf"
ui_ctx.input_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py\\7. Lernziele Herz_ Kreislauf und Gefasssystem.pdf"
# ui_ctx.input_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py\\Fragen.pdf"
# ui_ctx.input_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py\\Lernziele Bewegungsapparat.pdf"
# ui_ctx.input_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py\\FilesReading\\Fragen.pdf"
# ui_ctx.output_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py\\FilesReading"
ui_ctx.output_file_name = "C:\\Users\\Nicolas\\GitHub\\Automation_Py"

create_questions_file(ui_ctx)

'''
https://pymupdf.readthedocs.io/en/latest/app1.html
<page>
    <text block>
        <line>
            <span>
                <char>
    <image block>
        <img>
A text page consists of blocks (= roughly paragraphs).

A block consists of either lines and their characters, or an image.

A line consists of spans.

A span consists of adjacent characters with identical font properties: name, size, flags and color.

'''